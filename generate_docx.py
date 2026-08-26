import sys
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_docx():
    doc = Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    PRIMARY = RGBColor(45, 122, 58)      # Emerald green #2D7A3A
    SECONDARY = RGBColor(194, 94, 0)     # Amber accent #C25E00
    DARK = RGBColor(30, 41, 59)          # Slate #1E293B
    MUTED = RGBColor(100, 116, 139)      # Slate muted #64748B

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("🛒 KIRANA POINT")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Next-Generation Local Grocery E-Commerce & PWA Platform\nComprehensive Client Feature Proposal & Technical Capability Document")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = DARK

    # Meta Info Card Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Store Owner / Client:", "Mr. Pratham Tarde"),
        ("Store Name & Location:", "Kirana Point — Main Road, Khamgaon, Dist. Buldhana, Maharashtra 444303"),
        ("Direct UPI & Phone:", "8208232735@axl  |  +91 8208232735"),
        ("Platform Architecture:", "Next.js 14 (App Router), TypeScript, TailwindCSS, Firebase Firestore, PWA, Leaflet GPS"),
    ]
    for idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[idx]
        cell_0 = row.cells[0]
        cell_1 = row.cells[1]
        
        cell_0.width = Inches(2.2)
        cell_1.width = Inches(4.5)
        
        p0 = cell_0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = PRIMARY
        
        p1 = cell_1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.size = Pt(10)
        r1.font.color.rgb = DARK
        
        set_cell_background(cell_0, "F0FDF4")
        set_cell_background(cell_1, "F8FAFC")
        set_cell_margins(cell_0, top=100, bottom=100, left=150, right=150)
        set_cell_margins(cell_1, top=100, bottom=100, left=150, right=150)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 1. Executive Summary
    h1 = doc.add_heading("1. Executive Summary & Business Advantage", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.add_run(
        "Kirana Point is an enterprise-grade, lightning-fast digital grocery platform built specifically for "
        "neighbourhood grocery retail. It empowers store owners like Mr. Pratham Tarde to compete directly and win "
        "against corporate quick-commerce giants (Blinkit, Zepto, Swiggy Instamart) with complete technological independence."
    )

    # Key Comparison Table
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Metric / Capability", "Corporate Apps (Blinkit / Zepto)", "Kirana Point (Your Platform)"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "2D7A3A")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

    rows_content = [
        ("Platform Commission", "15% - 25% deducted per order", "0% Commission — 100% profit stays with store"),
        ("Payment Gateway Fees", "2% - 3% gateway deduction", "₹0 Gateway Fees — Direct UPI to 8208232735@axl"),
        ("Customer Data & Ownership", "Owned by third-party aggregator", "100% Owned by Pratham Tarde (Direct customer list)"),
        ("Delivery Pricing Logic", "Dynamic surge pricing & hidden fees", "Transparent distance-based (₹5/km, Free > ₹2,000)"),
        ("Customer Support", "Impersonal automated bot support", "1-Click Direct WhatsApp chat with store owner"),
    ]

    for row_idx, row_data in enumerate(rows_content, start=1):
        row = table.rows[row_idx]
        for col_idx, text in enumerate(row_data):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
            elif col_idx == 2:
                r.font.bold = True
                r.font.color.rgb = PRIMARY
            bg = "F0FDF4" if col_idx == 2 else ("F8FAFC" if row_idx % 2 == 0 else "FFFFFF")
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 2. Detailed Feature Breakdown
    h1 = doc.add_heading("2. Complete Feature Breakdown", level=1)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    features = [
        (
            "A. Customer Storefront & Mobile-First PWA Experience",
            [
                ("Guest-First Free Browsing:", "Customers can freely scroll products, view deals, browse categories, and check nutrition without any forced login barriers upfront."),
                ("Action-Triggered Authentication:", "When an unauthenticated visitor clicks 'Add to Cart' or 'Proceed to Checkout', a sleek Login / Register portal opens seamlessly. Once logged in, their pending action immediately executes!"),
                ("Mobile Dedicated Search Bar:", "Full-width, thumb-friendly search bar with real-time autocomplete suggestions across product names, brands, categories, and tags."),
                ("Progressive Web App (PWA):", "Customers can install Kirana Point directly onto their Android or iOS home screen without downloading from app stores."),
                ("Sticky Action Bar on Mobile:", "When browsing product details on mobile phones, a persistent bottom bar allows instant 1-tap cart addition without scrolling."),
            ]
        ),
        (
            "B. Zero-Fee Direct UPI Payment System",
            [
                ("Direct-to-Bank Transfers:", "All payments route directly into Pratham Tarde's UPI ID (8208232735@axl) with ₹0 middleman gateway deductions."),
                ("UPI Deep Linking (GPay / PhonePe / Paytm):", "Customers on mobile phones can tap Google Pay, PhonePe, or Paytm buttons which pre-fill the exact order amount and store payee name automatically."),
                ("Dynamic QR Code Generator:", "Desktop shoppers get an automatically generated UPI QR code that can be scanned using any mobile camera or scanner app."),
                ("Payment Claim & Verification Queue:", "Customers tap 'I Have Completed Payment' after transferring, which sends the order into the Admin Verification queue for 1-click confirmation."),
                ("Cash on Delivery (COD) Support:", "Customers also have the flexibility to select Cash on Delivery with 1-click."),
            ]
        ),
        (
            "C. Precision GPS Distance-Based Delivery Engine",
            [
                ("Haversine Distance Algorithm:", "Calculates precise road distance between the store hub in Khamgaon (20.6865° N, 76.5654° E) and the customer's delivery pin."),
                ("Transparent Pricing Rules:", "Free delivery automatically applies for orders of ₹2,000 or above, or store pickup. Orders below ₹2,000 are charged at an honest ₹5 per kilometer (min ₹20)."),
                ("Interactive Leaflet Map:", "Customers can drag the map marker to their exact doorstep or tap 'Use My GPS' for 1-click geolocation detection."),
                ("Service Radius Guard:", "Orders beyond the store's 15km delivery perimeter are prompted to choose store pickup, preventing unserviceable deliveries."),
            ]
        ),
        (
            "D. Open Food Facts Automated Cataloging & Competitor Pricing",
            [
                ("1-Click Barcode / Product Auto-Fill:", "Store managers can enter a barcode or search term to auto-fetch official descriptions, ingredients, package quantities, and full nutrition tables from the Open Food Facts global database."),
                ("Competitor Price Comparison Widget:", "Each product page features a live comparison showing Kirana Point's low price vs Blinkit, Zepto, BigBasket, and JioMart, highlighting customer savings!"),
            ]
        ),
        (
            "E. WhatsApp Business Automation",
            [
                ("1-Click Invoice Notification:", "Formatted WhatsApp invoice drafts with order breakdown, items list, total amount, and delivery address sent directly via wa.me."),
                ("Direct Customer Support:", "A floating WhatsApp button allows customers to connect with the store manager in 1 tap for custom grocery requests or special delivery notes."),
            ]
        ),
        (
            "F. Complete Store Management Console (Admin Portal)",
            [
                ("Visual Order Status Pipeline:", "Interactive stepper moving orders through: Confirmed ➔ Packing & Preparing ➔ Out for Delivery ➔ Delivered / Picked Up."),
                ("Payment Verification Panel:", "Dedicated queue highlighting unverified UPI transfers for quick verification."),
                ("Inventory & Low Stock Alerts:", "Real-time indicators warning when essential staples or dairy products drop below threshold."),
                ("Product Catalog Management:", "Full control to add new products, edit pricing/discounts, update images, or toggle stock availability."),
                ("Revenue Analytics & Reports:", "Visual charts tracking daily/weekly sales volume, popular items, and customer transaction counts."),
                ("Store Settings Customizer:", "Easily update delivery rates per km, free delivery thresholds, operating hours, and store coordinates."),
            ]
        ),
    ]

    for sec_title, bullet_items in features:
        h2 = doc.add_heading(sec_title, level=2)
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(4)
        for b_title, b_desc in bullet_items:
            bp = doc.add_paragraph(style='List Bullet')
            r_title = bp.add_run(b_title + " ")
            r_title.font.bold = True
            r_title.font.size = Pt(10)
            r_title.font.color.rgb = DARK
            r_desc = bp.add_run(b_desc)
            r_desc.font.size = Pt(10)
            r_desc.font.color.rgb = DARK

    # 3. Store Credentials & Access Guide
    h1 = doc.add_heading("3. Store Credentials & Access Guide", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)

    cred_table = doc.add_table(rows=3, cols=4)
    cred_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_headers = ["User Role", "Login Identifier", "Password", "Target Destination"]
    for i, h in enumerate(c_headers):
        cell = cred_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "2D7A3A")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

    cred_rows = [
        ("Store Owner (Admin)", "pratham@kiranapoint.com\n(or mobile: 8208232735)", "admin123", "/admin (Store Console)"),
        ("Customer Account", "Any customer email / phone\n(or register new account)", "Customer's chosen password", "/ (Grocery Storefront)"),
    ]

    for r_idx, (role, ident, pwd, dest) in enumerate(cred_rows, start=1):
        row = cred_table.rows[r_idx]
        for c_idx, val in enumerate([role, ident, pwd, dest]):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            if c_idx == 0:
                r.font.bold = True
            set_cell_background(cell, "F8FAFC" if r_idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)

    # 4. Next Steps
    h1 = doc.add_heading("4. Deployment & Next Steps", level=1)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(6)
    
    steps_p = doc.add_paragraph()
    steps_p.add_run(
        "1. Cloud Hosting: The application is pushed to GitHub (https://github.com/Neilarmstrong-00/Kirana-Point) and is ready for 1-click deployment on Vercel.\n"
        "2. Custom Domain: Can be mapped to any custom domain (e.g. www.kiranapoint.in or www.kiranapoint.com).\n"
        "3. Live Operations: Store manager Pratham Tarde can log in, verify UPI payments, dispatch local orders, and update stock daily with zero technical complexity!"
    )

    docx_path = "D:\\Project Files\\Client Projects\\Kirana Point\\Kirana_Point_Client_Feature_Presentation.docx"
    doc.save(docx_path)
    print(f"DOCX created successfully at: {docx_path}")
    return docx_path

if __name__ == "__main__":
    create_docx()
